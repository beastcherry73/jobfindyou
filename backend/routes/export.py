import re
import io
import json
import html
from flask import Blueprint, request, jsonify, send_file
from backend.decorators import login_required

export_bp = Blueprint("export", __name__)


def _section_key(title):
    """Classify a markdown section heading into a canonical section key.

    Case-insensitive and tolerant of harmless formatting differences
    (surrounding whitespace, punctuation, capitalization).
    """
    t = re.sub(r'[#*_\-:\s]+', ' ', title).strip().lower()
    if not t:
        return None
    if any(k in t for k in ('summary', 'profile', 'objective')):
        return 'summary'
    if 'education' in t:
        if any(k in t for k in ('certif', 'licens', 'credential')):
            return 'education_certs'
        return 'education'
    if any(k in t for k in ('certif', 'licens', 'credential')):
        return 'certifications'
    if 'project' in t:
        return 'projects'
    if any(k in t for k in ('skill', 'competenc', 'technolog', 'core')):
        return 'skills'
    if any(k in t for k in ('experience', 'employment', 'work', 'career')):
        return 'experience'
    return 'custom'


def _looks_like_cert(line):
    """Heuristic: does this education-ish line look like a certification instead?"""
    l = line.lower()
    if any(k in l for k in ('certif', 'licens', 'credential', 'cert.', 'badge', 'course')):
        return True
    if re.search(r'(aws|gcp|google|azure|microsoft|coursera|udemy|linkedin|pmp|comptia|scrum|cisco)\b.*(20\d\d|\d{2})', l):
        return True
    return False


def parse_markdown_to_data(text):
    """
    Utility to parse raw markdown text into a complete structured dictionary.

    Recognizes (case-insensitive): Summary / Professional Summary,
    Skills / Core Competencies / Technical Skills / Professional Skills,
    Experience / Work Experience / Professional Experience,
    Education, Certifications, Education & Certifications (split heuristically),
    Projects. Content lines that do not start with ###, -, *, or bullet
    characters are still captured so valid content is never silently dropped.
    """
    lines = text.split('\n')
    data = {
        "fullName": "",
        "email": "",
        "phone": "",
        "location": "",
        "summary": "",
        "skills": "",
        "experience": [],
        "education": [],
        "projects": [],
        "certifications": [],
        "customSections": [],
        "rawText": text
    }

    current_sec = None
    buf = []

    for line in lines:
        s = line.strip()
        if not s:
            continue
        if s.startswith('# '):
            data["fullName"] = s[2:].strip()
            continue
        if s.startswith('## '):
            sec_title = s[3:].strip()
            key = _section_key(sec_title)
            current_sec = key
            if key == 'custom':
                data["customSections"].append({"title": sec_title, "lines": []})
            continue
        if current_sec is None:
            continue
        if current_sec == 'summary':
            data["summary"] = (data["summary"] + " " + s).strip()
        elif current_sec == 'skills':
            clean_s = s.lstrip('-•* ').strip()
            data["skills"] = (data["skills"] + ", " + clean_s if data["skills"] else clean_s).strip(', ')
        elif current_sec == 'custom':
            if data["customSections"]:
                data["customSections"][-1]["lines"].append(s.lstrip('-•* ').strip())
        else:
            buf.append((current_sec, s))

    # Entry parsing for buffer lines
    cur_entry = None
    for sec, b in buf:
        if sec == 'experience':
            if b.startswith('### '):
                if cur_entry:
                    data["experience"].append(cur_entry)
                header = b[4:].strip()
                dates = ""
                if '(' in header and ')' in header:
                    parts = header.rsplit('(', 1)
                    header = parts[0].strip()
                    dates = parts[1].rstrip(')').strip()
                title_parts = header.split(' - ') if ' - ' in header else header.split(' | ')
                role = title_parts[0].strip()
                company = title_parts[1].strip() if len(title_parts) > 1 else ""
                cur_entry = {"role": role, "company": company, "dates": dates, "bullets": []}
            elif b.startswith(('- ', '• ', '* ')):
                bullet = b.lstrip('-•* ').strip()
                if cur_entry:
                    cur_entry["bullets"].append(bullet)
                else:
                    cur_entry = {"role": "Experience", "company": "", "dates": "", "bullets": [bullet]}
            elif cur_entry:
                if not cur_entry["bullets"] and not cur_entry["company"]:
                    cur_entry["company"] = b
                else:
                    cur_entry["bullets"].append(b)
            else:
                cur_entry = {"role": "Experience", "company": "", "dates": "", "bullets": [b]}
        elif sec == 'education':
            if b.startswith(('### ', '- ', '• ', '* ')):
                clean_b = b.lstrip('#-•* ').strip()
                data["education"].append({"school": clean_b, "degree": "", "dates": ""})
            elif not data["education"]:
                data["education"].append({"school": b, "degree": "", "dates": ""})
            else:
                data["education"][-1]["school"] = (data["education"][-1]["school"] + " " + b).strip()
        elif sec == 'education_certs':
            if b.startswith(('### ', '- ', '• ', '* ')):
                clean_b = b.lstrip('#-•* ').strip()
                if _looks_like_cert(clean_b):
                    data["certifications"].append({"name": clean_b, "issuer": "", "dates": ""})
                else:
                    data["education"].append({"school": clean_b, "degree": "", "dates": ""})
            elif not data["education"] and not data["certifications"]:
                data["education"].append({"school": b, "degree": "", "dates": ""})
            else:
                if _looks_like_cert(b) and data["certifications"]:
                    data["certifications"][-1]["name"] = (data["certifications"][-1]["name"] + " " + b).strip()
                elif data["education"]:
                    data["education"][-1]["school"] = (data["education"][-1]["school"] + " " + b).strip()
                else:
                    data["education"].append({"school": b, "degree": "", "dates": ""})
        elif sec == 'projects':
            if b.startswith('### '):
                p_name = b[4:].strip()
                data["projects"].append({"name": p_name, "desc": "", "tech": ""})
            elif b.startswith(('- ', '• ', '* ')) and data["projects"]:
                data["projects"][-1]["desc"] = (data["projects"][-1]["desc"] + " " + b.lstrip('-•* ').strip()).strip()
            elif data["projects"]:
                data["projects"][-1]["desc"] = (data["projects"][-1]["desc"] + " " + b).strip()
            else:
                data["projects"].append({"name": b, "desc": "", "tech": ""})
        elif sec == 'certifications':
            if b.startswith(('### ', '- ', '• ', '* ')):
                clean_b = b.lstrip('#-•* ').strip()
                data["certifications"].append({"name": clean_b, "issuer": "", "dates": ""})
            elif not data["certifications"]:
                data["certifications"].append({"name": b, "issuer": "", "dates": ""})
            else:
                data["certifications"][-1]["name"] = (data["certifications"][-1]["name"] + " " + b).strip()

    if cur_entry and cur_entry not in data["experience"]:
        data["experience"].append(cur_entry)

    return data


def merge_missing_sections(data, raw_text):
    """Recover missing structured fields section-by-section from rawText.

    A single missing field must not cause unrelated sections to disappear,
    and must not trigger the whole-resume raw-line fallback.
    """
    if not data:
        return data
    d = dict(data)
    source_text = raw_text or d.get("rawText") or ""
    if not source_text:
        return d
    parsed = parse_markdown_to_data(source_text)
    for key in ("summary", "skills", "experience", "education", "projects", "certifications", "fullName", "customSections"):
        if not d.get(key) and parsed.get(key):
            d[key] = parsed[key]
    return d


def esc(text):
    """Safely escape text for ReportLab XML paragraph tags."""
    if not text:
        return ""
    return html.escape(str(text))


@export_bp.route("/api/export/parse", methods=["POST"])
@login_required
def parse_resume_text():
    """Return the canonical structured resume parsed from raw markdown/text.

    Used by saveOptimizedResumeVersion() so an AI-enhanced resume is persisted
    as structured data (experience/education/skills/projects/certifications),
    NOT just {fullName, summary, rawText}.
    """
    req_json = request.get_json() or {}
    text = req_json.get("text", "") or req_json.get("resume", "")
    if not text or not str(text).strip():
        return jsonify({"error": "No text provided to parse"}), 400
    parsed = parse_markdown_to_data(str(text))
    parsed["rawText"] = str(text)
    return jsonify({"data": parsed})


@export_bp.route("/api/export/pdf", methods=["POST"])
@login_required
def export_pdf():
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    req_json = request.get_json() or {}
    raw_data = req_json.get("data")
    raw_text = req_json.get("text", "").strip()
    title = req_json.get("title", "Resume").strip()
    template_style = (req_json.get("template") or "modern").lower()

    if isinstance(raw_data, str):
        try:
            raw_data = json.loads(raw_data)
        except Exception:
            raw_data = None

    if not raw_data and raw_text:
        raw_data = parse_markdown_to_data(raw_text)

    if not raw_data and not raw_text:
        return jsonify({"error": "No resume data or text provided for PDF export"}), 400

    data = raw_data or {}
    fallback_text = data.get("rawText") or raw_text
    data = merge_missing_sections(data, fallback_text)
    full_name = data.get("fullName") or data.get("name") or "Professional Resume"
    email = data.get("email", "")
    phone = data.get("phone", "")
    location = data.get("location", "")
    summary = data.get("summary", "")
    skills = data.get("skills", "")
    experience = data.get("experience", [])
    education = data.get("education", [])
    projects = data.get("projects", [])
    certifications = data.get("certifications", [])
    custom_sections = data.get("customSections", [])

    # Theme colors
    color_palette = {
        "modern": {"primary": colors.HexColor("#1E3A5F"), "accent": colors.HexColor("#2563EB"), "text": colors.HexColor("#0F172A")},
        "executive": {"primary": colors.HexColor("#0F172A"), "accent": colors.HexColor("#475569"), "text": colors.HexColor("#1E293B")},
        "creative": {"primary": colors.HexColor("#6D28D9"), "accent": colors.HexColor("#7C3AED"), "text": colors.HexColor("#1E1B4B")},
        "minimalist": {"primary": colors.HexColor("#334155"), "accent": colors.HexColor("#64748B"), "text": colors.HexColor("#1E293B")}
    }
    theme = color_palette.get(template_style, color_palette["modern"])

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()

    name_style = ParagraphStyle(
        'HeaderName',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=theme["primary"],
        spaceAfter=4
    )

    contact_style = ParagraphStyle(
        'HeaderContact',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=12,
        textColor=colors.HexColor('#64748B'),
        spaceAfter=8
    )

    sec_heading_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=15,
        textColor=theme["primary"],
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=theme["text"],
        spaceAfter=6
    )

    bullet_style = ParagraphStyle(
        'BulletCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=theme["text"],
        leftIndent=12,
        firstLineIndent=-8,
        spaceAfter=2.5
    )

    item_title_style = ParagraphStyle(
        'ItemTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        leading=13,
        textColor=colors.HexColor('#1E293B')
    )

    item_sub_style = ParagraphStyle(
        'ItemSub',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#64748B'),
        alignment=2  # Right aligned
    )

    story = []

    # 1. Header
    story.append(Paragraph(esc(full_name), name_style))
    contact_parts = [esc(p) for p in [email, phone, location] if p]
    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), contact_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=theme["accent"], spaceAfter=8))

    # 2. Professional Summary
    if summary:
        story.append(Paragraph("PROFESSIONAL SUMMARY", sec_heading_style))
        story.append(Paragraph(esc(summary), body_style))

    # 3. Skills Matrix
    if skills:
        story.append(Paragraph("SKILLS & COMPETENCIES", sec_heading_style))
        skills_text = skills if isinstance(skills, str) else ", ".join(skills)
        story.append(Paragraph(esc(skills_text), body_style))

    # 4. Work Experience
    if experience:
        story.append(Paragraph("WORK EXPERIENCE", sec_heading_style))
        for exp in experience:
            if isinstance(exp, dict):
                role = exp.get("role") or exp.get("title") or "Position"
                company = exp.get("company", "")
                dates = exp.get("dates", "")
                bullets = exp.get("bullets", [])

                title_line = f"<b>{esc(role)}</b>" + (f" — {esc(company)}" if company else "")
                t_table = Table(
                    [[Paragraph(title_line, item_title_style), Paragraph(esc(dates), item_sub_style)]],
                    colWidths=[380, 160]
                )
                t_table.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('TOPPADDING', (0, 0), (-1, -1), 2),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                ]))
                story.append(t_table)

                for b in bullets:
                    if b and str(b).strip():
                        story.append(Paragraph(f"• {esc(b)}", bullet_style))
                story.append(Spacer(1, 4))

    # 5. Education
    if education:
        story.append(Paragraph("EDUCATION", sec_heading_style))
        for edu in education:
            if isinstance(edu, dict):
                degree = edu.get("degree", "")
                school = edu.get("school", "")
                dates = edu.get("dates", "")
                title_line = f"<b>{esc(degree)}</b>" + (f" — {esc(school)}" if school else "")
                t_table = Table(
                    [[Paragraph(title_line, item_title_style), Paragraph(esc(dates), item_sub_style)]],
                    colWidths=[380, 160]
                )
                t_table.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('TOPPADDING', (0, 0), (-1, -1), 2),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                ]))
                story.append(t_table)
                story.append(Spacer(1, 4))

    # 6. Projects
    if projects:
        story.append(Paragraph("PROJECTS", sec_heading_style))
        for proj in projects:
            if isinstance(proj, dict):
                p_name = proj.get("name") or proj.get("title") or "Project"
                p_desc = proj.get("desc") or proj.get("description") or ""
                p_tech = proj.get("tech") or proj.get("technologies") or ""
                story.append(Paragraph(f"<b>{esc(p_name)}</b>" + (f" ({esc(p_tech)})" if p_tech else ""), item_title_style))
                if p_desc:
                    story.append(Paragraph(esc(p_desc), body_style))
                story.append(Spacer(1, 4))

    # 7. Certifications
    if certifications:
        story.append(Paragraph("CERTIFICATIONS & LICENSES", sec_heading_style))
        for cert in certifications:
            if isinstance(cert, dict):
                c_name = cert.get("name") or cert.get("title") or "Certification"
                c_issuer = cert.get("issuer") or cert.get("organization") or ""
                c_dates = cert.get("dates") or cert.get("date") or ""
                title_line = f"<b>{esc(c_name)}</b>" + (f" — {esc(c_issuer)}" if c_issuer else "")
                t_table = Table(
                    [[Paragraph(title_line, item_title_style), Paragraph(esc(c_dates), item_sub_style)]],
                    colWidths=[380, 160]
                )
                t_table.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('TOPPADDING', (0, 0), (-1, -1), 2),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                ]))
                story.append(t_table)
                story.append(Spacer(1, 4))
            elif isinstance(cert, str) and cert.strip():
                story.append(Paragraph(f"• {esc(cert)}", bullet_style))

    # 8. Custom sections parsed from markdown (headings not otherwise classified)
    if custom_sections:
        for section in custom_sections:
            if isinstance(section, dict):
                title = section.get("title") or "Additional"
                lines = section.get("lines") or []
                story.append(Paragraph(esc(title).upper(), sec_heading_style))
                for cl in lines:
                    if cl and str(cl).strip():
                        story.append(Paragraph(f"• {esc(cl)}", bullet_style))

    # Fallback to lines parsing only when NO structured sections were populated.
    # Having a summary alone must not silently drop the rest of the resume.
    if fallback_text and not experience and not education and not projects and not certifications:
        for line in fallback_text.split('\n'):
            s = line.strip()
            if not s:
                continue
            if s.startswith('# '):
                story.append(Paragraph(esc(s[2:]), name_style))
            elif s.startswith('## '):
                story.append(Paragraph(esc(s[3:]), sec_heading_style))
            elif s.startswith('### '):
                story.append(Paragraph(esc(s[4:]), item_title_style))
            elif s.startswith(('- ', '• ', '* ')):
                story.append(Paragraph(f"• {esc(s.lstrip('-•* '))}", bullet_style))
            else:
                story.append(Paragraph(esc(s), body_style))

    doc.build(story)
    buffer.seek(0)
    clean_filename = re.sub(r'[^a-zA-Z0-9_-]', '_', title) + '.pdf'
    return send_file(
        buffer,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=clean_filename
    )


@export_bp.route("/api/export/docx", methods=["POST"])
@login_required
def export_docx():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    req_json = request.get_json() or {}
    raw_data = req_json.get("data")
    raw_text = req_json.get("text", "").strip()
    title = req_json.get("title", "Optimized_Resume").strip()

    if isinstance(raw_data, str):
        try:
            raw_data = json.loads(raw_data)
        except Exception:
            raw_data = None

    if not raw_data and raw_text:
        raw_data = parse_markdown_to_data(raw_text)

    doc = Document()

    # Set 0.6 inch margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(15, 23, 42)

    if raw_data and (raw_data.get("fullName") or raw_data.get("experience") or raw_data.get("summary") or raw_data.get("rawText")):
        data = merge_missing_sections(raw_data, raw_data.get("rawText") or raw_text)
        full_name = data.get("fullName") or data.get("name") or "Professional Resume"
        email = data.get("email", "")
        phone = data.get("phone", "")
        location = data.get("location", "")
        summary = data.get("summary", "")
        skills = data.get("skills", "")
        experience = data.get("experience", [])
        education = data.get("education", [])
        projects = data.get("projects", [])
        certifications = data.get("certifications", [])
        custom_sections = data.get("customSections", [])

        # Header
        p = doc.add_paragraph()
        run = p.add_run(full_name)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(37, 99, 235)
        p.paragraph_format.space_after = Pt(2)

        contact_parts = [cp for cp in [email, phone, location] if cp]
        if contact_parts:
            p_c = doc.add_paragraph()
            r_c = p_c.add_run(" | ".join(contact_parts))
            r_c.font.size = Pt(9.5)
            r_c.font.color.rgb = RGBColor(100, 116, 139)
            p_c.paragraph_format.space_after = Pt(10)

        # Summary
        if summary:
            p_h = doc.add_paragraph()
            r_h = p_h.add_run("PROFESSIONAL SUMMARY")
            r_h.font.size = Pt(12)
            r_h.font.bold = True
            r_h.font.color.rgb = RGBColor(30, 41, 59)
            p_h.paragraph_format.space_before = Pt(8)
            p_h.paragraph_format.space_after = Pt(3)

            p_s = doc.add_paragraph()
            p_s.add_run(summary)
            p_s.paragraph_format.space_after = Pt(8)

        # Skills
        if skills:
            p_h = doc.add_paragraph()
            r_h = p_h.add_run("SKILLS & COMPETENCIES")
            r_h.font.size = Pt(12)
            r_h.font.bold = True
            r_h.font.color.rgb = RGBColor(30, 41, 59)
            p_h.paragraph_format.space_before = Pt(8)
            p_h.paragraph_format.space_after = Pt(3)

            p_sk = doc.add_paragraph()
            p_sk.add_run(skills if isinstance(skills, str) else ", ".join(skills))
            p_sk.paragraph_format.space_after = Pt(8)

        # Experience
        if experience:
            p_h = doc.add_paragraph()
            r_h = p_h.add_run("WORK EXPERIENCE")
            r_h.font.size = Pt(12)
            r_h.font.bold = True
            r_h.font.color.rgb = RGBColor(30, 41, 59)
            p_h.paragraph_format.space_before = Pt(8)
            p_h.paragraph_format.space_after = Pt(3)

            for exp in experience:
                if isinstance(exp, dict):
                    role = exp.get("role") or exp.get("title") or "Position"
                    company = exp.get("company", "")
                    dates = exp.get("dates", "")
                    bullets = exp.get("bullets", [])

                    p_e = doc.add_paragraph()
                    r_r = p_e.add_run(role)
                    r_r.bold = True
                    if company:
                        p_e.add_run(f" — {company}")
                    if dates:
                        p_e.add_run(f"\t{dates}")
                    p_e.paragraph_format.space_after = Pt(2)

                    for b in bullets:
                        if b and str(b).strip():
                            p_b = doc.add_paragraph(style='List Bullet')
                            p_b.add_run(str(b))
                            p_b.paragraph_format.space_after = Pt(2)

        # Education
        if education:
            p_h = doc.add_paragraph()
            r_h = p_h.add_run("EDUCATION")
            r_h.font.size = Pt(12)
            r_h.font.bold = True
            r_h.font.color.rgb = RGBColor(30, 41, 59)
            p_h.paragraph_format.space_before = Pt(8)
            p_h.paragraph_format.space_after = Pt(3)

            for edu in education:
                if isinstance(edu, dict):
                    degree = edu.get("degree", "")
                    school = edu.get("school", "")
                    dates = edu.get("dates", "")
                    p_ed = doc.add_paragraph()
                    r_d = p_ed.add_run(degree)
                    r_d.bold = True
                    if school:
                        p_ed.add_run(f" — {school}")
                    if dates:
                        p_ed.add_run(f"\t{dates}")
                    p_ed.paragraph_format.space_after = Pt(2)

        # Projects
        if projects:
            p_h = doc.add_paragraph()
            r_h = p_h.add_run("PROJECTS")
            r_h.font.size = Pt(12)
            r_h.font.bold = True
            r_h.font.color.rgb = RGBColor(30, 41, 59)
            p_h.paragraph_format.space_before = Pt(8)
            p_h.paragraph_format.space_after = Pt(3)

            for proj in projects:
                if isinstance(proj, dict):
                    p_name = proj.get("name") or proj.get("title") or "Project"
                    p_desc = proj.get("desc") or proj.get("description") or ""
                    p_tech = proj.get("tech") or proj.get("technologies") or ""
                    p_pj = doc.add_paragraph()
                    r_pn = p_pj.add_run(p_name)
                    r_pn.bold = True
                    if p_tech:
                        p_pj.add_run(f" ({p_tech})")
                    p_pj.paragraph_format.space_after = Pt(2)
                    if p_desc:
                        p_pd = doc.add_paragraph()
                        p_pd.add_run(p_desc)
                        p_pd.paragraph_format.space_after = Pt(4)

        # Certifications
        if certifications:
            p_h = doc.add_paragraph()
            r_h = p_h.add_run("CERTIFICATIONS & LICENSES")
            r_h.font.size = Pt(12)
            r_h.font.bold = True
            r_h.font.color.rgb = RGBColor(30, 41, 59)
            p_h.paragraph_format.space_before = Pt(8)
            p_h.paragraph_format.space_after = Pt(3)

            for cert in certifications:
                if isinstance(cert, dict):
                    c_name = cert.get("name") or cert.get("title") or "Certification"
                    c_issuer = cert.get("issuer") or cert.get("organization") or ""
                    c_dates = cert.get("dates") or cert.get("date") or ""
                    p_ct = doc.add_paragraph()
                    r_cn = p_ct.add_run(c_name)
                    r_cn.bold = True
                    if c_issuer:
                        p_ct.add_run(f" — {c_issuer}")
                    if c_dates:
                        p_ct.add_run(f"\t{c_dates}")
                    p_ct.paragraph_format.space_after = Pt(2)
                elif isinstance(cert, str) and cert.strip():
                    p_b = doc.add_paragraph(style='List Bullet')
                    p_b.add_run(cert)
                    p_b.paragraph_format.space_after = Pt(2)

        # Custom sections
        if custom_sections:
            for section in custom_sections:
                if isinstance(section, dict):
                    title = section.get("title") or "Additional"
                    lines = section.get("lines") or []
                    p_h = doc.add_paragraph()
                    r_h = p_h.add_run(title.upper())
                    r_h.font.size = Pt(12)
                    r_h.font.bold = True
                    r_h.font.color.rgb = RGBColor(30, 41, 59)
                    p_h.paragraph_format.space_before = Pt(8)
                    p_h.paragraph_format.space_after = Pt(3)
                    for cl in lines:
                        if cl and str(cl).strip():
                            p_b = doc.add_paragraph(style='List Bullet')
                            p_b.add_run(str(cl))
                            p_b.paragraph_format.space_after = Pt(2)

    else:
        # Fallback to lines parsing
        lines = raw_text.split('\n')
        for line in lines:
            line_s = line.strip()
            if not line_s:
                continue
            if line_s.startswith('# '):
                p = doc.add_paragraph()
                run = p.add_run(line_s[2:].strip())
                run.font.size = Pt(22)
                run.font.bold = True
                run.font.color.rgb = RGBColor(37, 99, 235)
                p.paragraph_format.space_after = Pt(4)
            elif line_s.startswith('## '):
                p = doc.add_paragraph()
                run = p.add_run(line_s[3:].strip())
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(30, 41, 59)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
            elif line_s.startswith('### '):
                p = doc.add_paragraph()
                run = p.add_run(line_s[4:].strip())
                run.font.size = Pt(11)
                run.font.bold = True
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
            elif line_s.startswith(('• ', '- ', '* ')):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(line_s.lstrip('-•* ').strip())
                p.paragraph_format.space_after = Pt(2)
            else:
                p = doc.add_paragraph()
                p.add_run(line_s)
                p.paragraph_format.space_after = Pt(4)

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)

    clean_filename = re.sub(r'[^a-zA-Z0-9_-]', '_', title) + '.docx'
    return send_file(
        target_stream,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=clean_filename
    )
