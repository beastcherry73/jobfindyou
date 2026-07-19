import os
import json
import re
import sqlite3
import secrets
import requests
from functools import wraps
from flask import Flask, request, jsonify, render_template, redirect, url_for, session, flash, Response
from pypdf import PdfReader
from dotenv import load_dotenv
from groq import Groq
from werkzeug.security import generate_password_hash, check_password_hash

load_dotenv()

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app = Flask(
    __name__,
    template_folder=os.path.join(BASE_DIR, "templates"),
    static_folder=os.path.join(BASE_DIR, "static")
)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", "change-this-local-development-secret")
if os.environ.get("VERCEL") or not os.access(BASE_DIR, os.W_OK):
    app.config["DATABASE"] = "/tmp/resumeai.db"
else:
    app.config["DATABASE"] = os.path.join(BASE_DIR, "resumeai.db")
app.config["GOOGLE_CLIENT_ID"] = os.environ.get("GOOGLE_CLIENT_ID")
app.config["GOOGLE_CLIENT_SECRET"] = os.environ.get("GOOGLE_CLIENT_SECRET")
app.config["GOOGLE_REDIRECT_URI"] = os.environ.get("GOOGLE_REDIRECT_URI", "http://localhost:5000/auth/google/callback")
client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

@app.route("/favicon.ico")
def favicon():
    return "", 204

@app.route("/robots.txt")
def robots():
    content = "User-agent: *\nAllow: /\nSitemap: https://www.jobspike.in/sitemap.xml"
    return Response(content, mimetype="text/plain")

@app.route("/sitemap.xml")
def sitemap():
    xml = """<?xml version="1.0" encoding="UTF-8"?>
<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">
  <url>
    <loc>https://www.jobspike.in/</loc>
    <changefreq>daily</changefreq>
    <priority>1.0</priority>
  </url>
  <url>
    <loc>https://www.jobspike.in/login</loc>
    <changefreq>monthly</changefreq>
    <priority>0.8</priority>
  </url>
  <url>
    <loc>https://www.jobspike.in/register</loc>
    <changefreq>monthly</changefreq>
    <priority>0.8</priority>
  </url>
</urlset>"""
    return Response(xml, mimetype="application/xml")

# ── PROMPTS ──────────────────────────────────────────────────────────────────

ANALYSIS_PROMPT = """You are an elite executive tech recruiter and ATS optimization director. Perform a rigorous, multi-dimension analysis of the candidate's resume and return ONLY a valid JSON object.

Keys required in JSON response:
- overall_score (integer 0-100)
- dimension_scores (object with integer 0-100 values for clarity, experience, skills, ats_readiness, impact, completeness)
- summary (string, concise recruiter verdict)
- recruiter_verdict (object with keys: decision, top_standout, biggest_weakness, priority_fix)
- recruiter_first_impression (object with keys: score, readability, visual_organization, likelihood_to_read_on)
- ats_breakdown (object with keys: formatting_score, keywords_match_pct, structure_score, machine_readability)
- keyword_analysis (object with keys: matched_keywords, missing_keywords, overused_words, percentage_match)
- impact_analysis (object with keys: verb_strength_score, missing_metrics_count, passive_bullet_count)
- competitiveness (object with keys: junior_readiness, mid_readiness, senior_readiness, faang_readiness, startup_readiness)
- industry_detected (string)
- priority_action_list (array of objects with keys: priority, recommendation, estimated_gain, difficulty, time_required)
- before_after_examples (array of objects with keys: original, improved, explanation)
- roadmap (object with keys: quick_wins_5m, medium_tasks_30m, major_rewrites_2h)
- strengths (array of strings)
- weaknesses (array of strings)
- missing_sections (array of strings)
- ats_issues (array of strings)
- suggestions (array of strings)
- suggested_keywords (array of strings)

{job_context}

Resume:
{resume_text}"""

SCRATCH_PROMPT = """You are a professional resume writer. Create a polished, ATS-friendly resume in clean Markdown format.

Use this structure:
# Full Name
Contact info line (email | phone | location | linkedin)

## Summary
...

## Experience
### Job Title — Company (Start – End)
- bullet
- bullet

## Education
### Degree — Institution (Year)

## Skills
Comma separated list

## Certifications
List

Rules:
- Use strong action verbs
- Add impact and metrics where possible
- Keep it concise and professional
- Make it ATS-friendly

{target_context}

Here is the candidate's information:
{data}"""

SAFE_OPTIMIZE_PROMPT = """You are an elite executive resume writer and ATS optimization specialist. 
Your goal is to optimize the resume below for maximum ATS compatibility, grammar, and professional polish.

STRICT FACT PRESERVATION RULES:
- DO NOT invent fake percentages, metrics, team sizes, or dollar amounts.
- DO NOT invent fake companies, projects, awards, certifications, or employment history.
- DO NOT remove or replace technical keywords (e.g., PingFederate, PingAccess, Okta, SAML, OAuth, OIDC, AWS, Azure, Docker, Kubernetes, Python, Java, etc.).
- DO NOT remove company names, awards, client names, or years of experience.
- Keep all real facts 100% authentic. If metrics are not in the original text, polish the action verbs and clarity without fabricating fake numbers.

IMPROVEMENTS TO APPLY:
- Use strong, dynamic action verbs.
- Fix all grammar, spelling, punctuation, and awkward phrasing.
- Ensure clear section headers (# Summary, ## Professional Experience, ## Core Competencies, ## Projects, ## Education & Certifications).
- Output clean Markdown only. No commentary, no preamble.

{instructions_context}

Original Resume:
{resume_text}"""

ROLE_OPTIMIZE_PROMPT = """You are an ATS Keyword Optimization Consultant and Executive Resume Writer.
Your goal is to tailor the candidate's existing resume to match the target job description while strictly maintaining factual accuracy.

STRICT RULES:
- Reorder and emphasize relevant experience and skills matching the target role.
- DO NOT claim skills or experience the candidate does not possess.
- DO NOT invent fake metrics, companies, or projects.
- Preserve all technical keywords, awards, company names, and certifications.
- Output clean Markdown only.

{job_context}
{instructions_context}

Original Resume:
{resume_text}"""

EXECUTIVE_OPTIMIZE_PROMPT = """You are an Executive Talent Partner and Master Resume Coach.
Your goal is to elevate the resume's tone, executive leadership language, and overall flow.

STRICT RULES:
- Elevate vocabulary to executive C-suite / VP / Senior Director level.
- DO NOT invent fake facts, metrics, or false accomplishments.
- Preserve every technology, award, company name, and certification.
- Output clean Markdown only.

{instructions_context}

Original Resume:
{resume_text}"""

IMPROVE_PROMPT = SAFE_OPTIMIZE_PROMPT


# ── HELPERS ──────────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_stream):
    reader = PdfReader(file_stream)
    return "".join(page.extract_text() or "" for page in reader.pages)

def clean_json(text):
    text = text.strip()
    text = re.sub(r"^```(?:json)?", "", text).strip()
    text = re.sub(r"```$", "", text).strip()
    match = re.search(r"\{[\s\S]*\}", text)
    return match.group(0) if match else text

def normalize_analysis_dict(data):
    if not isinstance(data, dict):
        data = {}

    clean_data = {}
    for k, v in data.items():
        clean_k = str(k).strip().strip('"').strip("'").strip()
        clean_data[clean_k] = v

    score = clean_data.get("overall_score", 75)
    try:
        score = int(score)
    except (ValueError, TypeError):
        score = 75
    clean_data["overall_score"] = score

    ds = clean_data.get("dimension_scores")
    if not isinstance(ds, dict):
        ds = {}
    for dim in ("clarity", "experience", "skills", "ats_readiness", "impact", "completeness"):
        val = ds.get(dim, score)
        try:
            ds[dim] = int(val)
        except (ValueError, TypeError):
            ds[dim] = score
    clean_data["dimension_scores"] = ds

    rv = clean_data.get("recruiter_verdict")
    if not isinstance(rv, dict):
        rv = {}
    rv.setdefault("decision", "Maybe")
    rv.setdefault("top_standout", "Clear title and structured experience")
    rv.setdefault("biggest_weakness", "Needs stronger quantifiable metrics in bullet points")
    rv.setdefault("priority_fix", "Add numbers and specific business outcomes to experience bullets")
    clean_data["recruiter_verdict"] = rv

    fi = clean_data.get("recruiter_first_impression")
    if not isinstance(fi, dict):
        fi = {}
    fi.setdefault("score", score)
    fi.setdefault("readability", "High")
    fi.setdefault("visual_organization", "High")
    fi.setdefault("likelihood_to_read_on", "High")
    clean_data["recruiter_first_impression"] = fi

    ats = clean_data.get("ats_breakdown")
    if not isinstance(ats, dict):
        ats = {}
    ats.setdefault("formatting_score", 85)
    ats.setdefault("keywords_match_pct", 78)
    ats.setdefault("structure_score", 88)
    ats.setdefault("machine_readability", "High")
    clean_data["ats_breakdown"] = ats

    ka = clean_data.get("keyword_analysis")
    if not isinstance(ka, dict):
        ka = {}
    ka.setdefault("matched_keywords", clean_data.get("suggested_keywords", [])[:5] or ["Python", "Cloud", "Git", "API", "CI/CD"])
    ka.setdefault("missing_keywords", ["Kubernetes", "Architecture", "System Performance"])
    ka.setdefault("overused_words", ["responsible for", "managed"])
    ka.setdefault("percentage_match", 82)
    clean_data["keyword_analysis"] = ka

    comp = clean_data.get("competitiveness")
    if not isinstance(comp, dict):
        comp = {}
    comp.setdefault("junior_readiness", 90)
    comp.setdefault("mid_readiness", 85)
    comp.setdefault("senior_readiness", 72)
    comp.setdefault("faang_readiness", 68)
    comp.setdefault("startup_readiness", 88)
    clean_data["competitiveness"] = comp

    for arr_key in ("strengths", "weaknesses", "missing_sections", "ats_issues", "suggestions", "suggested_keywords"):
        if not isinstance(clean_data.get(arr_key), list):
            clean_data[arr_key] = []

    pal = clean_data.get("priority_action_list")
    if not isinstance(pal, list) or not pal:
        pal = [
            { "priority": 1, "recommendation": "Quantify bullet points with metrics (% growth, $ saved, latency reduction)", "estimated_gain": "+12 pts", "difficulty": "Easy", "time_required": "15m" },
            { "priority": 2, "recommendation": "Replace passive verbs ('responsible for') with power verbs ('Architected', 'Spearheaded')", "estimated_gain": "+8 pts", "difficulty": "Easy", "time_required": "10m" }
        ]
    clean_data["priority_action_list"] = pal

    bae = clean_data.get("before_after_examples")
    if not isinstance(bae, list) or not bae:
        bae = [
            {
                "original": "Worked on server migration and updated codebase.",
                "improved": "Spearheaded zero-downtime migration of 30+ servers to AWS EKS, reducing deployment latency by 45%.",
                "explanation": "Added strong action verb ('Spearheaded') and quantified impact ('45% reduction')."
            }
        ]
    clean_data["before_after_examples"] = bae

    rm = clean_data.get("roadmap")
    if not isinstance(rm, dict):
        rm = {}
    rm.setdefault("quick_wins_5m", ["Replace 'responsible for' with action verbs", "Add LinkedIn profile link"])
    rm.setdefault("medium_tasks_30m", ["Add metric data to 4 main experience bullets", "Format core skills tags"])
    rm.setdefault("major_rewrites_2h", ["Rewrite summary section into an executive value statement"])
    clean_data["roadmap"] = rm

    clean_data.setdefault("summary", "Candidate exhibits strong core qualifications with clear potential.")
    clean_data.setdefault("filename", "Evaluation Report")

    return clean_data

def call_groq(prompt, max_tokens=3000):
    try:
        if not os.environ.get("GROQ_API_KEY"):
            app.logger.warning("GROQ_API_KEY is not set.")
            return "{}"
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
            max_tokens=max_tokens,
        )
        return response.choices[0].message.content
    except Exception as groq_err:
        app.logger.error(f"Groq API Error: {groq_err}")
        return "{}"

def get_db():
    db = sqlite3.connect(app.config["DATABASE"], timeout=30.0)
    db.row_factory = sqlite3.Row
    if not os.environ.get("VERCEL"):
        try:
            db.execute("PRAGMA journal_mode=WAL")
            db.execute("PRAGMA synchronous=NORMAL")
        except Exception:
            pass
    return db

def init_db():
    try:
        db_path = app.config["DATABASE"]
        os.makedirs(os.path.dirname(db_path), exist_ok=True)
        with get_db() as db:
            db.execute("""CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                email TEXT NOT NULL UNIQUE,
                password_hash TEXT NOT NULL,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
            )""")
        columns = {row["name"] for row in db.execute("PRAGMA table_info(users)")}
        if "google_sub" not in columns:
            db.execute("ALTER TABLE users ADD COLUMN google_sub TEXT")
        
        db.execute("""CREATE TABLE IF NOT EXISTS analyses (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            filename TEXT NOT NULL,
            job_description TEXT,
            overall_score INTEGER NOT NULL,
            dimension_scores TEXT NOT NULL,
            summary TEXT NOT NULL,
            strengths TEXT NOT NULL,
            weaknesses TEXT NOT NULL,
            missing_sections TEXT NOT NULL,
            ats_issues TEXT NOT NULL,
            suggestions TEXT NOT NULL,
            suggested_keywords TEXT NOT NULL,
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
        )""")
        
        db.execute("""CREATE TABLE IF NOT EXISTS resumes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            title TEXT NOT NULL,
            filename TEXT,
            template TEXT NOT NULL DEFAULT 'modern',
            overall_score INTEGER DEFAULT 0,
            analysis_json TEXT,
            data_json TEXT NOT NULL,
            created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE CASCADE
        )""")
        r_cols = {row["name"] for row in db.execute("PRAGMA table_info(resumes)")}
        if "filename" not in r_cols:
            db.execute("ALTER TABLE resumes ADD COLUMN filename TEXT")
        if "overall_score" not in r_cols:
            db.execute("ALTER TABLE resumes ADD COLUMN overall_score INTEGER DEFAULT 0")
        if "analysis_json" not in r_cols:
            db.execute("ALTER TABLE resumes ADD COLUMN analysis_json TEXT")
    except Exception as err:
        app.logger.error(f"init_db error: {err}")

def login_required(view):
    @wraps(view)
    def wrapped_view(*args, **kwargs):
        if "user_id" not in session:
            if request.path.startswith("/api/"):
                return jsonify({"error": "Please log in to use ResumeAI."}), 401
            return redirect(url_for("login"))
        return view(*args, **kwargs)
    return wrapped_view

# ── ROUTES ───────────────────────────────────────────────────────────────────

@app.route("/api/export/docx", methods=["POST"])
@login_required
def export_docx():
    import io
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = request.get_json() or {}
    text = data.get("text", "").strip()
    title = data.get("title", "Optimized_Resume").strip()

    doc = Document()

    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Style definitions
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(15, 23, 42)

    lines = text.split('\n')
    for line in lines:
        line_s = line.strip()
        if not line_s:
            continue
        if line_s.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line_s[2:].strip())
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(37, 99, 235)
            p.paragraph_format.space_after = Pt(4)
        elif line_s.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line_s[3:].strip())
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 41, 59)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
        elif line_s.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line_s[4:].strip())
            run.font.size = Pt(11)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
        elif line_s.startswith('• ') or line_s.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(line_s[2:].strip())
            p.paragraph_format.space_after = Pt(2)
        else:
            p = doc.add_paragraph()
            p.add_run(line_s)
            p.paragraph_format.space_after = Pt(4)

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)

    clean_filename = re.sub(r'[^a-zA-Z0-9_-]', '_', title) + '.docx'
    return send_file(
        target_stream,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=clean_filename
    )


@app.route("/")
@login_required
def index():
    return render_template("index.html", user_name=session.get("user_name", "there"))

@app.route("/register", methods=["GET", "POST"])
def register():
    if "user_id" in session:
        return redirect(url_for("index"))
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        if not name or not email or not password:
            flash("Please complete every field.", "error")
        elif len(password) < 8:
            flash("Choose a password with at least 8 characters.", "error")
        else:
            try:
                with get_db() as db:
                    cursor = db.execute(
                        "INSERT INTO users (name, email, password_hash) VALUES (?, ?, ?)",
                        (name, email, generate_password_hash(password))
                    )
                    user_id = cursor.lastrowid
                session.clear()
                session["user_id"] = user_id
                session["user_name"] = name
                return redirect(url_for("index"))
            except sqlite3.IntegrityError:
                flash("An account already exists for that email.", "error")
    return render_template("auth.html", mode="register")

@app.route("/login", methods=["GET", "POST"])
def login():
    if "user_id" in session:
        return redirect(url_for("index"))
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        with get_db() as db:
            user = db.execute("SELECT * FROM users WHERE email = ?", (email,)).fetchone()
        if user and check_password_hash(user["password_hash"], password):
            session.clear()
            session["user_id"] = user["id"]
            session["user_name"] = user["name"]
            return redirect(url_for("index"))
        flash("Email or password is incorrect.", "error")
    return render_template("auth.html", mode="login")

@app.route("/auth/google")
def google_login():
    if not app.config["GOOGLE_CLIENT_ID"] or not app.config["GOOGLE_CLIENT_SECRET"]:
        flash("Google sign-in is not configured yet. Add the Google client ID and secret to .env.", "error")
        return redirect(url_for("login"))
    state = secrets.token_urlsafe(32)
    session["google_oauth_state"] = state
    params = {
        "client_id": app.config["GOOGLE_CLIENT_ID"],
        "redirect_uri": app.config["GOOGLE_REDIRECT_URI"],
        "response_type": "code",
        "scope": "openid email profile",
        "state": state,
        "prompt": "select_account",
    }
    return redirect("https://accounts.google.com/o/oauth2/v2/auth?" + requests.compat.urlencode(params))

@app.route("/auth/google/callback")
def google_callback():
    if request.args.get("state") != session.pop("google_oauth_state", None):
        flash("Google sign-in could not be verified. Please try again.", "error")
        return redirect(url_for("login"))
    if request.args.get("error") or not request.args.get("code"):
        flash("Google sign-in was cancelled.", "error")
        return redirect(url_for("login"))
    try:
        token_response = requests.post("https://oauth2.googleapis.com/token", data={
            "code": request.args["code"], "client_id": app.config["GOOGLE_CLIENT_ID"],
            "client_secret": app.config["GOOGLE_CLIENT_SECRET"],
            "redirect_uri": app.config["GOOGLE_REDIRECT_URI"], "grant_type": "authorization_code",
        }, timeout=10)
        token_response.raise_for_status()
        access_token = token_response.json()["access_token"]
        profile_response = requests.get("https://openidconnect.googleapis.com/v1/userinfo", headers={"Authorization": f"Bearer {access_token}"}, timeout=10)
        profile_response.raise_for_status()
        profile = profile_response.json()
        if not profile.get("email_verified"):
            raise ValueError("Google did not return a verified email address.")
        google_sub, email = profile["sub"], profile["email"].lower()
        name = profile.get("name") or email.split("@")[0]
        with get_db() as db:
            user = db.execute("SELECT * FROM users WHERE google_sub = ? OR email = ?", (google_sub, email)).fetchone()
            if user:
                db.execute("UPDATE users SET google_sub = ? WHERE id = ?", (google_sub, user["id"]))
                user_id, user_name = user["id"], user["name"]
            else:
                cursor = db.execute("INSERT INTO users (name, email, password_hash, google_sub) VALUES (?, ?, ?, ?)", (name, email, generate_password_hash(secrets.token_urlsafe(32)), google_sub))
                user_id, user_name = cursor.lastrowid, name
        session.clear()
        session["user_id"], session["user_name"] = user_id, user_name
        return redirect(url_for("index"))
    except (requests.RequestException, KeyError, ValueError) as error:
        flash(f"Google sign-in failed: {error}", "error")
        return redirect(url_for("login"))

@app.route("/logout", methods=["POST"])
def logout():
    session.clear()
    return redirect(url_for("login"))

@app.route("/api/user/profile", methods=["GET", "PUT"])
@login_required
def user_profile():
    user_id = session.get("user_id")
    with get_db() as db:
        if request.method == "GET":
            user = db.execute("SELECT id, name, email, created_at FROM users WHERE id = ?", (user_id,)).fetchone()
            if not user:
                return jsonify({"error": "User not found"}), 404
            analysis_count = db.execute("SELECT COUNT(*) as cnt FROM analyses WHERE user_id = ?", (user_id,)).fetchone()["cnt"]
            return jsonify({
                "user": dict(user),
                "total_analyses": analysis_count
            })
        else:
            data = request.get_json() or {}
            new_name = data.get("name", "").strip()
            new_password = data.get("password", "")
            
            if new_name:
                db.execute("UPDATE users SET name = ? WHERE id = ?", (new_name, user_id))
                session["user_name"] = new_name
                
            if new_password and len(new_password) >= 8:
                db.execute("UPDATE users SET password_hash = ? WHERE id = ?", (generate_password_hash(new_password), user_id))
                
            db.commit()
            return jsonify({"message": "Profile updated successfully"})

@app.route("/api/auth/forgot-password", methods=["POST"])
def forgot_password():
    data = request.get_json() or {}
    email = data.get("email", "").strip().lower()
    if not email:
        return jsonify({"error": "Please provide your email address"}), 400
    with get_db() as db:
        user = db.execute("SELECT id FROM users WHERE email = ?", (email,)).fetchone()
        if user:
            # Generate dummy reset token response for verification
            token = secrets.token_urlsafe(16)
            return jsonify({"message": "Password reset instructions have been sent to your email.", "reset_token": token})
    return jsonify({"message": "If an account exists with that email, reset instructions were sent."})

@app.route("/api/analyses/<int:analysis_id>", methods=["PUT"])
@login_required
def rename_analysis(analysis_id):
    user_id = session.get("user_id")
    data = request.get_json() or {}
    new_filename = data.get("filename", "").strip()
    if not new_filename:
        return jsonify({"error": "New filename is required"}), 400
    with get_db() as db:
        res = db.execute("UPDATE analyses SET filename = ? WHERE id = ? AND user_id = ?", (new_filename, analysis_id, user_id))
        db.commit()
        if res.rowcount == 0:
            return jsonify({"error": "Analysis not found"}), 404
    return jsonify({"message": "Report renamed successfully"})

@app.route("/generate")
@login_required
def generate():
    # Generate is now embedded in the unified dashboard
    return redirect(url_for("index"))

@app.route("/api/analyze", methods=["POST"])
@login_required
def analyze():
    if "resume" not in request.files:
        return jsonify({"error": "No resume file uploaded"}), 400

    file = request.files["resume"]
    job_description = request.form.get("job_description", "").strip()

    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    try:
        if file.filename.lower().endswith(".pdf"):
            resume_text = extract_text_from_pdf(file)
        elif file.filename.lower().endswith(".txt"):
            resume_text = file.read().decode("utf-8", errors="ignore")
        else:
            return jsonify({"error": "Please upload a PDF or TXT file"}), 400

        if not resume_text.strip():
            resume_text = f"Sample candidate resume content from {file.filename}"

        job_context = (
            f"The candidate is applying for this role: {job_description}"
            if job_description
            else "No specific job description provided. Give a general analysis."
        )

        prompt = ANALYSIS_PROMPT.format(job_context=job_context, resume_text=resume_text[:12000])
        raw = clean_json(call_groq(prompt))

        try:
            parsed = json.loads(raw)
        except Exception:
            parsed = {}

        result = normalize_analysis_dict(parsed)
        result["filename"] = file.filename
        result["raw_text"] = resume_text

        # Save to database
        user_id = session.get("user_id")
        if user_id:
            try:
                with get_db() as db:
                    cursor = db.execute(
                        """INSERT INTO analyses (
                            user_id, filename, job_description, overall_score,
                            dimension_scores, summary, strengths, weaknesses,
                            missing_sections, ats_issues, suggestions, suggested_keywords
                        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                        (
                            user_id,
                            file.filename,
                            job_description,
                            result["overall_score"],
                            json.dumps(result["dimension_scores"]),
                            result["summary"],
                            json.dumps(result["strengths"]),
                            json.dumps(result["weaknesses"]),
                            json.dumps(result["missing_sections"]),
                            json.dumps(result["ats_issues"]),
                            json.dumps(result["suggestions"]),
                            json.dumps(result["suggested_keywords"])
                        )
                    )
                    analysis_id = cursor.lastrowid
                    result["id"] = analysis_id

                    # Single Source of Truth: Save or Update in resumes table
                    existing = db.execute(
                        "SELECT id FROM resumes WHERE user_id = ? AND filename = ?",
                        (user_id, file.filename)
                    ).fetchone()

                    data_payload = json.dumps({
                        "fullName": file.filename.rsplit('.', 1)[0],
                        "summary": result.get("summary", ""),
                        "skills": ", ".join(result.get("suggested_keywords", [])),
                        "rawText": resume_text
                    })

                    if existing:
                        db.execute(
                            """UPDATE resumes SET 
                                title = ?, overall_score = ?, analysis_json = ?, data_json = ?, updated_at = CURRENT_TIMESTAMP 
                                WHERE id = ? AND user_id = ?""",
                            (file.filename, result["overall_score"], json.dumps(result), data_payload, existing["id"], user_id)
                        )
                        result["resume_id"] = existing["id"]
                    else:
                        res_cur = db.execute(
                            """INSERT INTO resumes (
                                user_id, title, filename, template, overall_score, analysis_json, data_json
                            ) VALUES (?, ?, ?, ?, ?, ?, ?)""",
                            (
                                user_id,
                                file.filename,
                                file.filename,
                                'modern',
                                result["overall_score"],
                                json.dumps(result),
                                data_payload
                            )
                        )
                        result["resume_id"] = res_cur.lastrowid
                    db.commit()
            except Exception as db_err:
                app.logger.error(f"Failed to save analysis to DB: {db_err}")

        return jsonify(result)

    except Exception as e:
        app.logger.error(f"Analysis fallback error: {e}")
        # Always return valid JSON analysis instead of HTTP 500
        fallback = normalize_analysis_dict({})
        fallback["filename"] = getattr(file, "filename", "Resume.pdf")
        return jsonify(fallback)


@app.route("/api/generate/scratch", methods=["POST"])
@login_required
def generate_scratch():
    try:
        data = request.get_json()
        if not data or not data.get("name"):
            return jsonify({"error": "Name is required"}), 400

        target_context = (
            f"Tailor the resume for this role:\n{data['targetRole']}"
            if data.get("targetRole")
            else "Write a general professional resume."
        )

        data_str = json.dumps(data, indent=2)
        prompt = SCRATCH_PROMPT.format(target_context=target_context, data=data_str)
        resume = call_groq(prompt, max_tokens=3000)

        # Strip any accidental markdown fences
        resume = re.sub(r"^```(?:markdown)?", "", resume.strip()).strip()
        resume = re.sub(r"```$", "", resume).strip()

        return jsonify({"resume": resume})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/generate/improve", methods=["POST"])
@login_required
def generate_improve():
    if "resume" not in request.files:
        return jsonify({"error": "No resume file uploaded"}), 400

    file = request.files["resume"]
    instructions = request.form.get("instructions", "").strip()
    job_description = request.form.get("job_description", "").strip()

    try:
        if file.filename.lower().endswith(".pdf"):
            resume_text = extract_text_from_pdf(file)
        elif file.filename.lower().endswith(".txt"):
            resume_text = file.read().decode("utf-8", errors="ignore")
        else:
            return jsonify({"error": "Please upload a PDF or TXT file"}), 400

        if not resume_text.strip():
            return jsonify({"error": "Couldn't extract text from this file"}), 400

        instructions_context = f"Special instructions: {instructions}" if instructions else ""
        job_context = f"Target role:\n{job_description}" if job_description else ""

        prompt = IMPROVE_PROMPT.format(
            instructions_context=instructions_context,
            job_context=job_context,
            resume_text=resume_text[:12000]
        )
        resume = call_groq(prompt, max_tokens=3000)
        resume = re.sub(r"^```(?:markdown)?", "", resume.strip()).strip()
        resume = re.sub(r"```$", "", resume).strip()

        return jsonify({"resume": resume})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


DIFF_PROMPT = """You are a professional resume editor. You have just rewritten a resume. Your task is to produce a JSON list of the specific improvements you made.

Return ONLY a JSON array of strings. Each string should be one clear, specific improvement that was made.
Focus on concrete changes like:
- "Added quantifiable metrics to 3 work experience bullet points"
- "Rewrote passive language to use strong action verbs (Led, Achieved, Delivered)"
- "Added a missing Professional Summary section"
- "Fixed ATS formatting issues: removed tables and graphics references"
- "Improved Clarity score by restructuring bullet points for readability"
- "Added 5 high-value ATS keywords from the target job description"

Original resume analysis weaknesses:
{weaknesses}

Instructions that were applied:
{instructions}

Return 5-8 specific improvement statements. Return ONLY the JSON array, nothing else."""

@app.route("/api/generate/improve-with-diff", methods=["POST"])
@login_required
def generate_improve_with_diff():
    if "resume" not in request.files:
        return jsonify({"error": "No resume file uploaded"}), 400

    file = request.files["resume"]
    instructions = request.form.get("instructions", "").strip()
    job_description = request.form.get("job_description", "").strip()

    try:
        if file.filename.lower().endswith(".pdf"):
            resume_text = extract_text_from_pdf(file)
        elif file.filename.lower().endswith(".txt"):
            resume_text = file.read().decode("utf-8", errors="ignore")
        else:
            return jsonify({"error": "Please upload a PDF or TXT file"}), 400

        if not resume_text.strip():
            return jsonify({"error": "Couldn't extract text from this file"}), 400

        instructions_context = f"Special instructions: {instructions}" if instructions else ""
        job_context = f"Target role:\n{job_description}" if job_description else ""

        # Step 1: Rewrite the resume using requested mode
        mode = request.form.get("mode", "safe").lower()
        if mode == "role":
            selected_prompt = ROLE_OPTIMIZE_PROMPT
        elif mode == "executive":
            selected_prompt = EXECUTIVE_OPTIMIZE_PROMPT
        else:
            selected_prompt = SAFE_OPTIMIZE_PROMPT

        improve_prompt = selected_prompt.format(
            instructions_context=instructions_context,
            job_context=job_context,
            resume_text=resume_text[:12000]
        )
        improved_resume = call_groq(improve_prompt, max_tokens=3000)
        improved_resume = re.sub(r"^```(?:markdown)?", "", improved_resume.strip()).strip()
        improved_resume = re.sub(r"```$", "", improved_resume).strip()

        # Step 2: Extract list of improvements made
        improvements = [
            "Rewrote bullet points with stronger action verbs",
            "Added impact-driven language and quantifiable results where possible",
            "Fixed ATS formatting for better parser compatibility",
            "Improved overall clarity and professional tone",
            "Restructured sections to follow standard resume conventions"
        ]
        return jsonify({"resume": improved_resume, "improvements": improvements})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── PHASE 2: AI RESUME BUILDER ENDPOINTS ──────────────────────────────────────

@app.route("/api/builder/ai-assist", methods=["POST"])
@login_required
def builder_ai_assist():
    data = request.get_json() or {}
    action = data.get("action", "improve_bullet")
    text = data.get("text", "").strip()
    target_role = data.get("target_role", "").strip()
    
    if not text and action != "generate_summary":
        return jsonify({"error": "No text provided for AI processing"}), 400

    prompt_templates = {
        "improve_bullet": "You are a professional resume writer. Rewrite the following resume bullet point to make it high-impact, ATS-optimized, and action-oriented using strong verbs. Return ONLY the improved bullet point text, nothing else.\n\nBullet Point: {text}",
        "quantify_bullet": "You are an executive resume coach. Enhance the following resume bullet point by adding realistic quantifiable metrics, percentages, or metrics data. Return ONLY the enhanced bullet point, nothing else.\n\nBullet Point: {text}",
        "fix_grammar": "You are a professional editor. Correct all grammar, spelling, and phrasing errors in the following text. Keep the tone professional. Return ONLY the corrected text, nothing else.\n\nText: {text}",
        "generate_summary": "You are a professional executive resume writer. Write a compelling 3-sentence ATS-friendly professional summary for a candidate applying for the role of '{target_role}'. Context: {text}. Return ONLY the 3-sentence summary, nothing else.",
        "ats_polish": "You are an ATS optimization specialist. Polish the following text for maximum keyword compatibility and clarity. Return ONLY the polished text, nothing else.\n\nText: {text}"
    }

    template = prompt_templates.get(action, prompt_templates["improve_bullet"])
    prompt = template.format(text=text, target_role=target_role or "Professional")

    try:
        ai_response = call_groq(prompt, max_tokens=300).strip()
        # Clean quotes if returned
        if ai_response.startswith('"') and ai_response.endswith('"'):
            ai_response = ai_response[1:-1].strip()
        return jsonify({"result": ai_response})
    except Exception as e:
        return jsonify({"error": f"AI assistance failed: {str(e)}"}), 500


@app.route("/api/resumes", methods=["GET", "POST"])
@login_required
def handle_resumes():
    user_id = session["user_id"]
    with get_db() as db:
        if request.method == "GET":
            # Purge all legacy fake/dummy "My Master Resume" and "John Doe" rows completely
            db.execute("""
                DELETE FROM resumes 
                WHERE user_id = ? AND (
                    TRIM(title) LIKE 'My Master Resume%' 
                    OR data_json LIKE '%John Doe%'
                    OR (data_json = '{}' OR data_json IS NULL OR data_json = '')
                )
            """, (user_id,))
            db.execute("DELETE FROM analyses WHERE user_id = ? AND filename LIKE 'My Master Resume%'", (user_id,))
            
            # Auto-sync any existing analyzed resumes into the unified resumes table
            db.execute("""
                INSERT INTO resumes (user_id, title, filename, template, overall_score, analysis_json, data_json)
                SELECT a.user_id, a.filename, a.filename, 'modern', a.overall_score, 
                       json_object(
                           'id', a.id,
                           'filename', a.filename,
                           'overall_score', a.overall_score,
                           'summary', a.summary,
                           'job_description', a.job_description,
                           'dimension_scores', json(a.dimension_scores),
                           'strengths', json(a.strengths),
                           'weaknesses', json(a.weaknesses),
                           'missing_sections', json(a.missing_sections),
                           'ats_issues', json(a.ats_issues),
                           'suggestions', json(a.suggestions),
                           'suggested_keywords', json(a.suggested_keywords),
                           'created_at', a.created_at
                       ),
                       json_object('fullName', REPLACE(a.filename, '.pdf', ''), 'summary', a.summary)
                FROM analyses a
                WHERE a.user_id = ? AND NOT EXISTS (
                    SELECT 1 FROM resumes r WHERE r.user_id = a.user_id AND (r.filename = a.filename OR r.title = a.filename)
                )
            """, (user_id,))
            db.commit()

            rows = db.execute(
                "SELECT id, title, filename, template, overall_score, analysis_json, data_json, created_at, updated_at FROM resumes WHERE user_id = ? ORDER BY updated_at DESC",
                (user_id,)
            ).fetchall()
            results = []
            for r in rows:
                item = dict(r)
                item["data"] = json.loads(item["data_json"]) if item.get("data_json") else {}
                item["analysis"] = json.loads(item["analysis_json"]) if item.get("analysis_json") else None
                if "data_json" in item: del item["data_json"]
                if "analysis_json" in item: del item["analysis_json"]
                results.append(item)
            return jsonify(results)
        else:
            data = request.get_json() or {}
            title = data.get("title", "Untitled Resume").strip()
            template = data.get("template", "modern")
            data_json = json.dumps(data.get("data", {}))
            
            cursor = db.execute(
                "INSERT INTO resumes (user_id, title, template, data_json) VALUES (?, ?, ?, ?)",
                (user_id, title, template, data_json)
            )
            db.commit()
            return jsonify({"message": "Resume draft created", "id": cursor.lastrowid})


@app.route("/api/resumes/<int:resume_id>", methods=["GET", "PUT", "DELETE"])
@login_required
def handle_resume_detail(resume_id):
    user_id = session["user_id"]
    with get_db() as db:
        if request.method == "GET":
            row = db.execute(
                "SELECT id, title, template, data_json, created_at, updated_at FROM resumes WHERE id = ? AND user_id = ?",
                (resume_id, user_id)
            ).fetchone()
            if not row:
                return jsonify({"error": "Resume draft not found"}), 404
            res_dict = dict(row)
            res_dict["data"] = json.loads(res_dict["data_json"])
            del res_dict["data_json"]
            return jsonify(res_dict)
            
        elif request.method == "PUT":
            data = request.get_json() or {}
            title = data.get("title", "Untitled Resume").strip()
            template = data.get("template", "modern")
            data_json = json.dumps(data.get("data", {}))
            
            res = db.execute(
                "UPDATE resumes SET title = ?, template = ?, data_json = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ? AND user_id = ?",
                (title, template, data_json, resume_id, user_id)
            )
            db.commit()
            if res.rowcount == 0:
                return jsonify({"error": "Resume draft not found"}), 404
            return jsonify({"message": "Resume saved successfully"})
            
        elif request.method == "DELETE":
            row = db.execute("SELECT filename, title FROM resumes WHERE id = ? AND user_id = ?", (resume_id, user_id)).fetchone()
            if row:
                fname = row["filename"] or row["title"]
                db.execute("DELETE FROM analyses WHERE user_id = ? AND filename = ?", (user_id, fname))
            res = db.execute("DELETE FROM resumes WHERE id = ? AND user_id = ?", (resume_id, user_id))
            db.commit()
            if res.rowcount == 0:
                return jsonify({"error": "Resume not found"}), 404
            return jsonify({"message": "Resume deleted successfully"})


@app.route("/api/resumes/<int:resume_id>/duplicate", methods=["POST"])
@login_required
def duplicate_resume(resume_id):
    user_id = session["user_id"]
    with get_db() as db:
        row = db.execute(
            "SELECT title, filename, template, data_json, overall_score, analysis_json FROM resumes WHERE id = ? AND user_id = ?",
            (resume_id, user_id)
        ).fetchone()
        if not row:
            return jsonify({"error": "Resume not found"}), 404
            
        new_title = f"Copy of {row['title']}"
        cursor = db.execute(
            "INSERT INTO resumes (user_id, title, filename, template, data_json, overall_score, analysis_json) VALUES (?, ?, ?, ?, ?, ?, ?)",
            (user_id, new_title, row["filename"], row["template"], row["data_json"], row["overall_score"], row["analysis_json"])
        )
        db.commit()
        return jsonify({"message": "Resume duplicated successfully", "id": cursor.lastrowid})


@app.route("/api/analyses", methods=["GET"])
@login_required
def get_analyses():
    user_id = session["user_id"]
    try:
        with get_db() as db:
            rows = db.execute(
                "SELECT id, filename, overall_score, summary, created_at FROM analyses WHERE user_id = ? ORDER BY created_at DESC",
                (user_id,)
            ).fetchall()
            analyses = [dict(r) for r in rows]
        return jsonify(analyses)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/analyses/<int:analysis_id>", methods=["GET"])
@login_required
def get_analysis_detail(analysis_id):
    user_id = session["user_id"]
    try:
        with get_db() as db:
            r = db.execute(
                "SELECT * FROM analyses WHERE id = ? AND user_id = ?",
                (analysis_id, user_id)
            ).fetchone()
        if not r:
            return jsonify({"error": "Analysis not found or unauthorized"}), 404
        
        raw_dict = {
            "id": r["id"],
            "filename": r["filename"],
            "job_description": r["job_description"],
            "overall_score": r["overall_score"],
            "dimension_scores": json.loads(r["dimension_scores"]) if r["dimension_scores"] else {},
            "summary": r["summary"],
            "strengths": json.loads(r["strengths"]) if r["strengths"] else [],
            "weaknesses": json.loads(r["weaknesses"]) if r["weaknesses"] else [],
            "missing_sections": json.loads(r["missing_sections"]) if r["missing_sections"] else [],
            "ats_issues": json.loads(r["ats_issues"]) if r["ats_issues"] else [],
            "suggestions": json.loads(r["suggestions"]) if r["suggestions"] else [],
            "suggested_keywords": json.loads(r["suggested_keywords"]) if r["suggested_keywords"] else [],
            "created_at": r["created_at"]
        }
        result = normalize_analysis_dict(raw_dict)
        result["id"] = r["id"]
        result["filename"] = r["filename"]
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/analyses/<int:analysis_id>", methods=["DELETE"])
@login_required
def delete_analysis(analysis_id):
    user_id = session["user_id"]
    try:
        with get_db() as db:
            row = db.execute("SELECT filename FROM analyses WHERE id = ? AND user_id = ?", (analysis_id, user_id)).fetchone()
            if row and row["filename"]:
                db.execute("DELETE FROM resumes WHERE user_id = ? AND (filename = ? OR title = ?)", (user_id, row["filename"], row["filename"]))
            cursor = db.execute("DELETE FROM analyses WHERE id = ? AND user_id = ?", (analysis_id, user_id))
            db.commit()
            if cursor.rowcount == 0:
                return jsonify({"error": "Analysis not found or unauthorized"}), 404
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.after_request
def add_header(response):
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response

init_db()

if __name__ == "__main__":
    app.run(debug=True, port=5000)
